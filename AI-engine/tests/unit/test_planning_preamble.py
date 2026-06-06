from __future__ import annotations

import io

from docx import Document

from app.application.services.planning_docx_extractor import (
    PLAN_INFORMATION_SECTION_ID,
    analyze_docx_template,
)
from app.application.services.planning_template_structure import document_template_from_analysis


def test_preamble_becomes_plan_information_section() -> None:
    doc = Document()
    doc.add_paragraph("Client-facing plan template | Prepared for RM review and client discussion")
    doc.add_paragraph("Plan Information Value / Placeholder")
    doc.add_paragraph("Client / Household [Client name / household name]")
    doc.add_paragraph("Relationship Manager / Advisor [RM / advisor name]")
    doc.add_paragraph("Plan date [DD/MM/YYYY]")
    doc.add_paragraph("Important note: This plan is prepared for discussion and planning purposes.")
    doc.add_paragraph("1. Executive Summary")
    doc.add_paragraph("Primary objective [Client objective and target date]")
    buf = io.BytesIO()
    doc.save(buf)

    analysis = analyze_docx_template(buf.getvalue())
    logical = analysis.get("logicalSections", [])
    assert logical
    assert logical[0]["id"] == PLAN_INFORMATION_SECTION_ID
    slots = logical[0].get("bracketSlots") or []
    assert any("Client name" in s for s in slots)
    assert any("RM / advisor" in s for s in slots)

    template = document_template_from_analysis(analysis, locale="en-US")
    assert template["sections"][0]["id"] == PLAN_INFORMATION_SECTION_ID
