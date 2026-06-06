from __future__ import annotations

import base64
import io
import re
from typing import Any

from docx import Document


def _normalize_tag(key: str) -> str:
    trimmed = (key or "").strip()
    if not trimmed:
        return trimmed
    if trimmed.startswith("{{") and trimmed.endswith("}}"):
        return trimmed
    if trimmed.startswith("{{"):
        return trimmed + "}}"
    if trimmed.endswith("}}"):
        return "{{" + trimmed
    return "{{" + trimmed + "}}"


def _alias_keys(tag: str) -> list[str]:
    """Match template tags with or without braces."""
    norm = _normalize_tag(tag)
    inner = norm[2:-2] if norm.startswith("{{") and norm.endswith("}}") else norm
    return list(
        dict.fromkeys(
            [
                norm,
                inner,
                inner.upper(),
                inner.lower(),
                f"«{inner}»",
                f"${{{inner}}}",
            ]
        )
    )


def _apply_replacements(text: str, replacements: dict[str, str]) -> str:
    updated = text
    for key, value in replacements.items():
        if not key or value is None:
            continue
        for alias in _alias_keys(key):
            if alias in updated:
                updated = updated.replace(alias, value)
    return updated


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> bool:
    runs = list(paragraph.runs)
    full_text = "".join(run.text for run in runs) if runs else (paragraph.text or "")
    if not full_text.strip():
        return False
    updated = _apply_replacements(full_text, replacements)
    if updated == full_text:
        return False
    if runs:
        runs[0].text = updated
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def _walk_paragraphs(document: Document, replacements: dict[str, str]) -> int:
    replaced = 0
    for paragraph in document.paragraphs:
        if _replace_in_paragraph(paragraph, replacements):
            replaced += 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _replace_in_paragraph(paragraph, replacements):
                        replaced += 1
    for section in document.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            if _replace_in_paragraph(paragraph, replacements):
                replaced += 1
        for paragraph in footer.paragraphs:
            if _replace_in_paragraph(paragraph, replacements):
                replaced += 1
    return replaced


def _section_export_order(section: dict[str, Any]) -> tuple[int, str]:
    sid = str(section.get("id") or "")
    if sid == "plan_information" or section.get("preamble"):
        return (0, "")
    number = str(section.get("number") or "").strip()
    if number.isdigit():
        return (int(number), str(section.get("title") or ""))
    return (999, str(section.get("title") or ""))


def _iter_document_sections(
    document_template: dict[str, Any] | None,
    section_content: dict[str, str] | None,
) -> list[dict[str, Any]]:
    if not isinstance(section_content, dict) or not section_content:
        return []
    sections: list[dict[str, Any]] = []
    if isinstance(document_template, dict):
        raw = document_template.get("sections")
        if isinstance(raw, list):
            sections = [s for s in raw if isinstance(s, dict)]
    if not sections:
        sections = [{"id": k, "title": k, "include": True} for k in section_content.keys()]
    return sorted(sections, key=_section_export_order)


def _write_llm_document_sections(
    document: Document,
    document_template: dict[str, Any] | None,
    section_content: dict[str, str] | None,
    *,
    page_break_before: bool = False,
) -> bool:
    """Write per-section LLM bodies into the document."""
    sections = _iter_document_sections(document_template, section_content)
    if not sections:
        return False

    written = False
    if page_break_before:
        document.add_page_break()
    for section in sections:
        if section.get("include") is False:
            continue
        section_id = str(section.get("id") or "").strip()
        if not section_id:
            continue
        body = str(section_content.get(section_id) or "").strip()
        if not body:
            continue
        heading = str(section.get("title") or section_id).strip()
        document.add_heading(heading, level=2)
        for block in re.split(r"\n\s*\n", body):
            line = block.strip()
            if line:
                document.add_paragraph(line)
        written = True
    return written


def _append_llm_document_sections(
    document: Document,
    document_template: dict[str, Any] | None,
    section_content: dict[str, str] | None,
    *,
    title: str = "Nội dung kế hoạch (AI)",
) -> bool:
    """Append LLM sections after an existing template document."""
    if not isinstance(section_content, dict) or not section_content:
        return False
    document.add_page_break()
    document.add_heading(title, level=1)
    return _write_llm_document_sections(
        document,
        document_template,
        section_content,
        page_break_before=False,
    )


def _write_metadata_lines(document: Document, placeholders: dict[str, str]) -> None:
    lines: list[str] = []
    for key, label in (
        ("{{CLIENT_ID}}", "Mã khách hàng"),
        ("{{CASE_ID}}", "Mã case"),
        ("{{TEMPLATE_CODE}}", "Mẫu kế hoạch"),
        ("{{GENERATED_AT}}", "Ngày lập"),
    ):
        value = placeholders.get(key) or placeholders.get(key.replace("{{", "").replace("}}", ""))
        if value:
            lines.append(f"{label}: {value}")
    for line in lines:
        document.add_paragraph(line)


def build_llm_only_docx(
    *,
    document_template: dict[str, Any] | None,
    section_content: dict[str, str] | None,
    placeholders: dict[str, str] | None = None,
    narratives: dict[str, Any] | None = None,
    title: str | None = None,
) -> tuple[bytes, int]:
    """Create a fresh Word document containing only LLM-composed plan content."""
    document = Document()
    plan_title = title
    if not plan_title and isinstance(document_template, dict):
        locale = str(document_template.get("locale") or "").strip()
        plan_title = "Kế hoạch tài chính" + (f" ({locale})" if locale else "")
    if not plan_title:
        plan_title = "Kế hoạch tài chính"
    document.add_heading(plan_title, level=0)
    if placeholders:
        _write_metadata_lines(document, placeholders)

    hits = 0
    if _write_llm_document_sections(document, document_template, section_content, page_break_before=False):
        hits += 1
    elif narratives:
        _append_ai_summary(document, narratives, title="Tóm tắt")
        hits += 1

    out = io.BytesIO()
    document.save(out)
    return out.getvalue(), hits


def _append_ai_summary(document: Document, narratives: dict[str, Any], *, title: str) -> None:
    exec_sum = str(narratives.get("executiveSummary") or "").strip()
    situation = str(narratives.get("situationAnalysis") or "").strip()
    recommendations = str(narratives.get("recommendations") or "").strip()
    data_quality = str(narratives.get("dataQualityNotes") or "").strip()
    if not any([exec_sum, situation, recommendations, data_quality]):
        return

    document.add_page_break()
    document.add_heading(title, level=1)

    sections = [
        ("Tóm lược điều hành", exec_sum),
        ("Phân tích hiện trạng", situation),
        ("Khuyến nghị", recommendations),
        ("Ghi chú chất lượng dữ liệu", data_quality),
    ]
    for heading, body in sections:
        if not body:
            continue
        document.add_heading(heading, level=2)
        for block in re.split(r"\n\s*\n", body):
            line = block.strip()
            if line:
                document.add_paragraph(line)


def _resolve_export_mode(export_mode: str, section_content: dict[str, str] | None) -> str:
    mode = (export_mode or "auto").strip().lower()
    if mode == "auto":
        return "llm_only" if section_content else "merge_template"
    if mode in {"llm_only", "merge_template"}:
        return mode
    return "llm_only" if section_content else "merge_template"


def render_docx(
    *,
    template_bytes: bytes,
    placeholders: dict[str, str],
    narratives: dict[str, Any] | None = None,
    document_template: dict[str, Any] | None = None,
    section_content: dict[str, str] | None = None,
    export_mode: str = "auto",
    append_summary: bool = True,
    append_document_sections: bool = False,
    summary_title: str = "Tóm tắt kế hoạch (AI)",
    document_title: str | None = None,
) -> tuple[bytes, int]:
    """
    Export planning DOCX.

    - ``llm_only``: fresh document from LLM ``sectionContent`` (no uploaded template body).
    - ``merge_template``: fill placeholders in the uploaded DOCX only.
    - ``auto``: ``llm_only`` when section content exists, else ``merge_template``.
    """
    sections = section_content if isinstance(section_content, dict) else {}
    mode = _resolve_export_mode(export_mode, sections)

    if mode == "llm_only":
        return build_llm_only_docx(
            document_template=document_template,
            section_content=sections,
            placeholders=placeholders,
            narratives=narratives,
            title=document_title,
        )

    document = Document(io.BytesIO(template_bytes))
    hits = _walk_paragraphs(document, placeholders)

    if append_document_sections:
        if _append_llm_document_sections(document, document_template, sections, title=summary_title):
            hits += 1
    elif append_summary and narratives:
        _append_ai_summary(document, narratives, title=summary_title)
        hits += 1

    out = io.BytesIO()
    document.save(out)
    return out.getvalue(), hits


def render_docx_base64(
    *,
    template_base64: str,
    placeholders: dict[str, str],
    narratives: dict[str, Any] | None = None,
    document_template: dict[str, Any] | None = None,
    section_content: dict[str, str] | None = None,
    export_mode: str = "auto",
    append_summary: bool = True,
    append_document_sections: bool = False,
    summary_title: str = "Tóm tắt kế hoạch (AI)",
    document_title: str | None = None,
) -> dict[str, Any]:
    sections = section_content if isinstance(section_content, dict) else {}
    mode = _resolve_export_mode(export_mode, sections)
    if mode == "llm_only":
        docx_bytes, hits = build_llm_only_docx(
            document_template=document_template,
            section_content=sections,
            placeholders=placeholders,
            narratives=narratives,
            title=document_title,
        )
    else:
        raw = base64.b64decode(template_base64) if template_base64 else b""
        if not raw:
            raise ValueError("templateBase64 is required for merge_template export mode.")
        docx_bytes, hits = render_docx(
            template_bytes=raw,
            placeholders=placeholders,
            narratives=narratives,
            document_template=document_template,
            section_content=sections,
            export_mode="merge_template",
            append_summary=append_summary,
            append_document_sections=append_document_sections,
            summary_title=summary_title,
            document_title=document_title,
        )
    return {
        "docxBase64": base64.b64encode(docx_bytes).decode("ascii"),
        "exportMode": mode,
        "replacementHits": hits,
        "byteSize": len(docx_bytes),
    }
