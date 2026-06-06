from __future__ import annotations

import io

from docx import Document

from app.application.services.planning_docx_renderer import build_llm_only_docx, render_docx


def test_render_docx_merge_template_replaces_placeholders() -> None:
    base = Document()
    base.add_paragraph("Hello {{EXECUTIVE_SUMMARY}}")
    buf = io.BytesIO()
    base.save(buf)
    template_bytes = buf.getvalue()

    docx_bytes, hits = render_docx(
        template_bytes=template_bytes,
        placeholders={"{{EXECUTIVE_SUMMARY}}": "Tóm lược từ LLM"},
        narratives={},
        document_template={},
        section_content={},
        export_mode="merge_template",
        append_summary=False,
    )
    out = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in out.paragraphs)
    assert "Tóm lược từ LLM" in full_text
    assert "Hello" not in full_text or "Tóm lược từ LLM" in full_text
    assert hits >= 1


def test_build_llm_only_docx_without_template_body() -> None:
    docx_bytes, hits = build_llm_only_docx(
        document_template={
            "sections": [
                {"id": "executive", "title": "Tóm lược điều hành", "include": True},
                {"id": "situation", "title": "Phân tích hiện trạng", "include": True},
            ]
        },
        section_content={
            "executive": "Đoạn tóm lược chi tiết do LLM sinh.",
            "situation": "Phân tích tài sản và dòng tiền.",
        },
        placeholders={"{{CLIENT_ID}}": "C-001"},
        title="Kế hoạch mẫu A",
    )
    out = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in out.paragraphs)
    assert "Kế hoạch mẫu A" in full_text
    assert "Đoạn tóm lược chi tiết do LLM sinh." in full_text
    assert "Phân tích tài sản và dòng tiền." in full_text
    assert "Mã khách hàng: C-001" in full_text
    assert "Nội dung kế hoạch (AI)" not in full_text
    assert hits >= 1


def test_render_docx_auto_uses_llm_only_when_sections_exist() -> None:
    base = Document()
    base.add_paragraph("Template gốc không được giữ")
    buf = io.BytesIO()
    base.save(buf)

    docx_bytes, _ = render_docx(
        template_bytes=buf.getvalue(),
        placeholders={},
        section_content={"intro": "Nội dung chỉ từ LLM."},
        export_mode="auto",
    )
    out = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in out.paragraphs)
    assert "Nội dung chỉ từ LLM." in full_text
    assert "Template gốc không được giữ" not in full_text
