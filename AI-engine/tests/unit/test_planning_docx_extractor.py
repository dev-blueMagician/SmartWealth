from __future__ import annotations

import io

from docx import Document

from app.application.services.planning_docx_extractor import analyze_docx_template


def test_analyze_docx_detects_placeholders() -> None:
    doc = Document()
    doc.add_paragraph("Khách hàng: {{CLIENT_NAME}}")
    doc.add_paragraph("Thu nhập: {{MONTHLY_INCOME}}")
    buf = io.BytesIO()
    doc.save(buf)

    analysis = analyze_docx_template(buf.getvalue())
    detected = analysis.get("detectedPlaceholders", [])
    assert "{{CLIENT_NAME}}" in detected
    assert "{{MONTHLY_INCOME}}" in detected
    assert analysis.get("blockCount", 0) >= 1


def test_analyze_docx_detects_bracket_slots() -> None:
    doc = Document()
    doc.add_paragraph("Client: [Client name / household name]")
    buf = io.BytesIO()
    doc.save(buf)

    analysis = analyze_docx_template(buf.getvalue())
    detected = analysis.get("detectedPlaceholders", [])
    assert "[Client name / household name]" in detected
