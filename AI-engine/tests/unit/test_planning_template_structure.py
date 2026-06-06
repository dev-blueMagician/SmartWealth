from __future__ import annotations

import io

from docx import Document

from app.application.services.planning_docx_extractor import analyze_docx_template
from app.application.services.planning_template_structure import (
    document_template_from_analysis,
    ensure_section_content_coverage,
    merge_document_templates,
)


def test_logical_sections_from_numbered_headings() -> None:
    doc = Document()
    doc.add_paragraph("1. Executive Summary")
    doc.add_paragraph("Primary objective [Client objective and target date]")
    doc.add_paragraph("2. Client and Household Profile")
    doc.add_paragraph("Client / household members [Names and relationships included in scope]")
    buf = io.BytesIO()
    doc.save(buf)

    analysis = analyze_docx_template(buf.getvalue())
    logical = analysis.get("logicalSections", [])
    assert len(logical) >= 2
    assert logical[0]["title"].startswith("Executive")
    slots = logical[0].get("bracketSlots") or []
    assert any("Client objective" in s for s in slots)

    template = document_template_from_analysis(analysis, locale="en-US")
    assert len(template["sections"]) >= 2
    assert template["sections"][0]["placeholders"]


def test_ensure_section_content_coverage_fills_gaps() -> None:
    document_template = {
        "sections": [
            {"id": "sec_1", "title": "One", "include": True, "draftHint": "Fill one"},
            {"id": "sec_2", "title": "Two", "include": True},
        ]
    }
    out = ensure_section_content_coverage(document_template, {"sec_1": "Done"}, data_quality_notes="Missing mortgage rate")
    assert out["sec_1"] == "Done"
    assert "sec_2" in out
    assert "bổ sung" in out["sec_2"].lower() or "RM" in out["sec_2"]


def test_merge_document_templates_preserves_base_order() -> None:
    base = {
        "sections": [
            {"id": "a", "title": "A", "placeholders": ["[x]"], "include": True},
            {"id": "b", "title": "B", "include": True},
        ]
    }
    overlay = {
        "sections": [
            {"id": "b", "title": "B updated", "draftHint": "hint"},
            {"id": "c", "title": "C", "include": True},
        ]
    }
    merged = merge_document_templates(base, overlay)
    ids = [s["id"] for s in merged["sections"]]
    assert ids == ["plan_information", "a", "b", "c"]
    assert merged["sections"][2]["title"] == "B updated"
