from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

_PLACEHOLDER_PATTERNS = (
    re.compile(r"\{\{([^}]{1,120})\}\}"),
    re.compile(r"«([^»]{1,120})»"),
    re.compile(r"\$\{([^}]{1,120})\}"),
)
# RM templates often use bracket slots e.g. [Client name / household name]
_BRACKET_SLOT_PATTERN = re.compile(r"\[([^\[\]]{1,200})\]")
_SECTION_HEADING_PATTERN = re.compile(r"^(\d{1,2})\.\s+(.+)$")
PLAN_INFORMATION_SECTION_ID = "plan_information"
PLAN_INFORMATION_TITLE = "Plan Information"
_DEFAULT_PLAN_INFO_BRACKET_SLOTS = [
    "[Client name / household name]",
    "[RM / advisor name]",
    "[DD/MM/YYYY]",
    "[Start date]",
    "[target date]",
    "[Save for home / Pay back home loan / Refinance / Hybrid plan]",
    "[VND / USD / other]",
    "[Draft / Reviewed / Final]",
]


def _normalize_placeholder_token(raw: str) -> str:
    inner = (raw or "").strip()
    if not inner:
        return ""
    if inner.startswith("{{") and inner.endswith("}}"):
        return inner
    return "{{" + inner + "}}"


def _extract_bracket_slots_from_text(text: str) -> list[str]:
    found: list[str] = []
    for match in _BRACKET_SLOT_PATTERN.finditer(text or ""):
        inner = match.group(1).strip()
        if not inner:
            continue
        slot = f"[{inner}]"
        if slot not in found:
            found.append(slot)
    return found


def _extract_placeholders_from_text(text: str) -> list[str]:
    found: list[str] = []
    for pattern in _PLACEHOLDER_PATTERNS:
        for match in pattern.finditer(text or ""):
            token = match.group(1).strip()
            norm = _normalize_placeholder_token(token)
            if norm and norm not in found:
                found.append(norm)
    for slot in _extract_bracket_slots_from_text(text):
        if slot not in found:
            found.append(slot)
    return found


def _section_id_from_title(title: str, number: str = "") -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", (title or "").lower()).strip("_")
    if number:
        return f"sec_{number}_{slug[:48]}" if slug else f"sec_{number}"
    return slug[:60] or "section"


def _new_section(*, number: str, title: str, preamble: bool = False) -> dict[str, Any]:
    return {
        "id": PLAN_INFORMATION_SECTION_ID if preamble else _section_id_from_title(title, number),
        "number": number,
        "title": PLAN_INFORMATION_TITLE if preamble else title,
        "placeholders": [],
        "bracketSlots": [],
        "previewLines": [],
        "purpose": (
            "Cover, plan metadata, client/RM/plan dates, disclaimer before numbered sections"
            if preamble
            else (f"Section {number}: {title}" if number else title)
        ),
        "preamble": preamble,
        "sortOrder": 0 if preamble else (int(number) * 10 if number.isdigit() else 500),
        "include": True,
    }


def _absorb_block_into_section(section: dict[str, Any], block: dict[str, Any]) -> None:
    text = str(block.get("text") or "").strip()
    if text and text not in section["previewLines"] and len(section["previewLines"]) < 30:
        section["previewLines"].append(text[:500])

    for tag in block.get("placeholders") or []:
        token = str(tag).strip()
        if not token:
            continue
        if token.startswith("[") and token.endswith("]"):
            if token not in section["bracketSlots"]:
                section["bracketSlots"].append(token)
        elif token not in section["placeholders"]:
            section["placeholders"].append(token)

    if block.get("type") == "table":
        for row in block.get("rowPreview") or []:
            if not isinstance(row, list):
                continue
            row_text = " ".join(str(c) for c in row)
            if row_text and row_text not in section["previewLines"]:
                section["previewLines"].append(row_text[:500])
            for slot in _extract_bracket_slots_from_text(row_text):
                if slot not in section["bracketSlots"]:
                    section["bracketSlots"].append(slot)
            for tag in _extract_placeholders_from_text(row_text):
                if tag.startswith("[") and tag.endswith("]"):
                    if tag not in section["bracketSlots"]:
                        section["bracketSlots"].append(tag)
                elif tag not in section["placeholders"]:
                    section["placeholders"].append(tag)
    elif text:
        for slot in _extract_bracket_slots_from_text(text):
            if slot not in section["bracketSlots"]:
                section["bracketSlots"].append(slot)


def _preamble_section_from_blocks(preamble_blocks: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not preamble_blocks:
        return None
    section = _new_section(number="", title=PLAN_INFORMATION_TITLE, preamble=True)
    for block in preamble_blocks:
        _absorb_block_into_section(section, block)
    for slot in _DEFAULT_PLAN_INFO_BRACKET_SLOTS:
        if slot not in section["bracketSlots"]:
            section["bracketSlots"].append(slot)
    return section if section.get("bracketSlots") or section.get("previewLines") else None


def _build_logical_sections(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Group blocks: preamble (before 1.) → plan_information; then numbered sections."""
    sections: list[dict[str, Any]] = []
    preamble_blocks: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    in_numbered = False

    def flush_current() -> None:
        nonlocal current
        if current and (current.get("placeholders") or current.get("bracketSlots") or current.get("title")):
            sections.append(current)
        current = None

    for block in blocks:
        text = str(block.get("text") or "").strip()
        heading_match = _SECTION_HEADING_PATTERN.match(text)
        is_numbered_start = heading_match is not None and bool(heading_match.group(1))
        is_heading = block.get("type") == "heading" or heading_match is not None

        if not in_numbered:
            if is_numbered_start and text:
                preamble = _preamble_section_from_blocks(preamble_blocks)
                if preamble:
                    sections.append(preamble)
                preamble_blocks = []
                in_numbered = True
                flush_current()
                number, title = heading_match.group(1), heading_match.group(2).strip()
                current = _new_section(number=number, title=title)
                continue
            preamble_blocks.append(block)
            continue

        if is_heading and text:
            flush_current()
            if heading_match:
                number, title = heading_match.group(1), heading_match.group(2).strip()
            else:
                number, title = "", text
            current = _new_section(number=number, title=title)
            continue

        if current is None:
            continue
        _absorb_block_into_section(current, block)

    if not in_numbered and preamble_blocks:
        preamble = _preamble_section_from_blocks(preamble_blocks)
        if preamble:
            sections.append(preamble)
    flush_current()
    return sections


def _collect_body_blocks_in_order(document: Document) -> tuple[list[dict[str, Any]], list[str]]:
    """Walk DOCX body in document order (paragraphs and tables interleaved)."""
    from docx.oxml.ns import qn

    blocks: list[dict[str, Any]] = []
    detected: list[str] = []
    idx = 0
    for element in document.element.body:
        if element.tag == qn("w:p"):
            rec = _paragraph_record(Paragraph(element, document), location="body", index=idx)
            if rec["text"] or rec["placeholders"]:
                blocks.append(rec)
                for tag in rec["placeholders"]:
                    if tag not in detected:
                        detected.append(tag)
            idx += 1
        elif element.tag == qn("w:tbl"):
            table = Table(element, document)
            table_tags: list[str] = []
            preview_rows: list[list[str]] = []
            for row in table.rows[:12]:
                cells: list[str] = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip().replace("\n", " ")
                    cells.append(cell_text[:200])
                    for tag in _extract_placeholders_from_text(cell_text):
                        if tag not in table_tags:
                            table_tags.append(tag)
                if any(cells):
                    preview_rows.append(cells)
            for tag in table_tags:
                if tag not in detected:
                    detected.append(tag)
            if preview_rows or table_tags:
                blocks.append(
                    {
                        "id": f"t-{idx}",
                        "type": "table",
                        "location": "body",
                        "rowPreview": preview_rows,
                        "placeholders": table_tags,
                    }
                )
            idx += 1
    return blocks, detected


def _paragraph_record(paragraph: Paragraph, *, location: str, index: int) -> dict[str, Any]:
    text = (paragraph.text or "").strip()
    style_name = paragraph.style.name if paragraph.style is not None else None
    placeholders = _extract_placeholders_from_text(text)
    heading_level: int | None = None
    if style_name and style_name.startswith("Heading"):
        try:
            heading_level = int(style_name.replace("Heading", "").strip())
        except ValueError:
            heading_level = None
    return {
        "id": f"p-{index}",
        "type": "heading" if heading_level else "paragraph",
        "location": location,
        "headingLevel": heading_level,
        "style": style_name,
        "text": text[:2000],
        "placeholders": placeholders,
    }


def _walk_paragraphs(document: Document, location: str, start_index: int) -> tuple[list[dict[str, Any]], list[str], int]:
    blocks: list[dict[str, Any]] = []
    all_tags: list[str] = []
    idx = start_index
    for paragraph in document.paragraphs:
        rec = _paragraph_record(paragraph, location=location, index=idx)
        if rec["text"] or rec["placeholders"]:
            blocks.append(rec)
            for tag in rec["placeholders"]:
                if tag not in all_tags:
                    all_tags.append(tag)
        idx += 1
    return blocks, all_tags, idx


def _walk_tables(document: Document, start_index: int) -> tuple[list[dict[str, Any]], list[str], int]:
    blocks: list[dict[str, Any]] = []
    all_tags: list[str] = []
    idx = start_index
    for table_index, table in enumerate(document.tables):
        if not isinstance(table, Table):
            continue
        table_tags: list[str] = []
        preview_rows: list[list[str]] = []
        for row in table.rows[:8]:
            cells: list[str] = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip().replace("\n", " ")
                cells.append(cell_text[:200])
                for tag in _extract_placeholders_from_text(cell_text):
                    if tag not in table_tags:
                        table_tags.append(tag)
            if any(cells):
                preview_rows.append(cells)
        for tag in table_tags:
            if tag not in all_tags:
                all_tags.append(tag)
        if preview_rows or table_tags:
            blocks.append(
                {
                    "id": f"t-{table_index}",
                    "type": "table",
                    "location": "body",
                    "rowPreview": preview_rows,
                    "placeholders": table_tags,
                }
            )
        idx += 1
    return blocks, all_tags, idx


def analyze_docx_template(template_bytes: bytes) -> dict[str, Any]:
    """
  Deterministic DOCX structure extraction for planning agent T1 (analyze_template).
  """
    document = Document(io.BytesIO(template_bytes))
    blocks, detected = _collect_body_blocks_in_order(document)
    idx = len(blocks)

    for section_index, section in enumerate(document.sections):
        header = section.header
        footer = section.footer
        if header is not None:
            h_blocks, h_tags, idx = _walk_paragraphs(header, f"header-{section_index}", idx)
            blocks.extend(h_blocks)
            for tag in h_tags:
                if tag not in detected:
                    detected.append(tag)
        if footer is not None:
            f_blocks, f_tags, idx = _walk_paragraphs(footer, f"footer-{section_index}", idx)
            blocks.extend(f_blocks)
            for tag in f_tags:
                if tag not in detected:
                    detected.append(tag)

    logical_sections = _build_logical_sections(blocks)
    return {
        "version": 1,
        "blockCount": len(blocks),
        "sections": blocks[:120],
        "logicalSections": logical_sections,
        "logicalSectionCount": len(logical_sections),
        "detectedPlaceholders": detected,
    }
